import requests
import shutil
import docx
import pandas as pd
from docx.shared import Inches

folder = r'C:\Users\DavidHolroyd\PycharmProjects\Pic_Downloader\Excel Input\\'
file = input('Please enter file name to coil picture file in Excel Input folder: ')
filepath = folder + file
coil_df = pd.read_excel(filepath, usecols=['Company', 'Name', 'Bldg', 'Floor', 'Rm', 'Pic1', 'Pic2',
                                           'Pic3', 'Pic4', 'Pic5', 'Pic6'])

coils = coil_df.values.tolist()  # This creates a list containing a list for each coil. This will be iterated over

doc = docx.Document()
doc.add_heading(f'{coils[0][0]} - Coil Pictures', 0)
coils_downloaded = 0
print(len(coil_df.columns))
num_coils = len(coils)

for coil in coils:
    coils_downloaded += 1
    print(f"Downloading pics: Coil #{coils_downloaded}/{num_coils}")
    pic_list = []
    for p in range(5, len(coil_df.columns)):
        if str(coil[p]) == 'nan':  # The picture columns contain a valid picture link, or nothing (output from app)
            pass
        else:
            pic_list.append(coil[p])

    # For each AHU coil, a collapsable header will be added with the format: Building / Floor / Room / Coil Name
    doc.add_heading(f'{coil[1]} / {coil[2]} / {coil[3]} / {coil[4]}', 3)

    for pic in pic_list:
        picture_obj = requests.get(fr'{pic}', stream=True)
        picture_obj.raw.decode_content = True
        with open("pic_saveover.jpg", 'wb') as f:  # Each picture is saved in the local directory
            shutil.copyfileobj(picture_obj.raw, f)
        doc.add_picture('pic_saveover.jpg', height=Inches(2.7))  # Saved picture is added to the Word doc
doc.save(fr'C:\Users\DavidHolroyd\PycharmProjects\Pic_Downloader\Word Docs\{coils[0][0]} Coil Pics.docx')
